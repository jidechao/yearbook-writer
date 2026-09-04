#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""将年鉴初稿文本转换为规范排版的 .docx。

输入文本格式约定：
    # 篇名标题            —— 首行，仅一个，居中加粗
    【条目标题】正文……   —— 条目，一段成文，首行缩进两字符
    （撰稿人）            —— 独立成行的括号署名，右对齐
    其余非空行            —— 普通正文段（首行缩进）

用法：
    python build_docx.py 输入.txt 输出.docx [--font-body 仿宋_GB2312]
        [--font-title 宋体] [--size 16] [--title-size 22]
"""
import argparse
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt


def set_run_font(run, font_name, size_pt, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # 中文字体需同时设置 eastAsia
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def add_paragraph(doc, text, font, size, bold=False, align=None, indent_chars=0):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    if indent_chars:
        # 用字符单位设置首行缩进
        pPr = p._p.get_or_add_pPr()
        ind = pPr.get_or_add_ind()
        ind.set(qn("w:firstLineChars"), str(indent_chars * 100))
    run = p.add_run(text)
    set_run_font(run, font, size, bold)
    return p


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("input", help="年鉴初稿文本文件（UTF-8）")
    ap.add_argument("output", help="输出 .docx 路径")
    ap.add_argument("--font-body", default="仿宋_GB2312")
    ap.add_argument("--font-title", default="宋体")
    ap.add_argument("--size", type=float, default=16, help="正文字号（磅），默认16=三号")
    ap.add_argument("--title-size", type=float, default=22, help="标题字号（磅），默认22=二号")
    args = ap.parse_args()

    with open(args.input, encoding="utf-8") as f:
        lines = [ln.rstrip() for ln in f]

    doc = Document()
    attr_re = re.compile(r"^（[^（）]{1,12}）$")

    for ln in lines:
        text = ln.strip()
        if not text:
            continue
        if text.startswith("# "):
            add_paragraph(doc, text[2:].strip(), args.font_title, args.title_size,
                          bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        elif attr_re.match(text) and not text.startswith("【"):
            add_paragraph(doc, text, args.font_body, args.size,
                          align=WD_ALIGN_PARAGRAPH.RIGHT)
        elif text.startswith("【"):
            # 条目标题加粗，正文常规
            m = re.match(r"^(【[^】]*】)(.*)$", text, re.S)
            head, rest = (m.group(1), m.group(2)) if m else ("", text)
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            ind = pPr.get_or_add_ind()
            ind.set(qn("w:firstLineChars"), "200")
            if head:
                r1 = p.add_run(head)
                set_run_font(r1, args.font_body, args.size, bold=True)
            if rest:
                r2 = p.add_run(rest)
                set_run_font(r2, args.font_body, args.size)
        else:
            add_paragraph(doc, text, args.font_body, args.size, indent_chars=2)

    doc.save(args.output)
    print(f"已生成：{args.output}")


if __name__ == "__main__":
    sys.exit(main())
